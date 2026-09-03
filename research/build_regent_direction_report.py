from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path
import re

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "regent-breakthrough-direction-2026-08-29.md"
OUT = ROOT / "Regent突破方向深度调研-2026-08-29.docx"

BLUE = "2E5B88"
DARK = "18324A"
MUTED = "667085"
LIGHT = "EEF3F8"
RISK = "9B1C1C"

def set_font(run, name="Microsoft YaHei", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa)); tc_w.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells): set_cell_width(cell, widths[i])

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), BLUE); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    rfonts = OxmlElement("w:rFonts"); rfonts.set(qn("w:eastAsia"), "Microsoft YaHei"); rfonts.set(qn("w:ascii"), "Microsoft YaHei"); rpr.append(rfonts)
    run.append(rpr); t = OxmlElement("w:t"); t.text = text; run.append(t); link.append(run); paragraph._p.append(link)

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for name, size, color, before, after in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",11.5,DARK,8,4)]:
    st = styles[name]; st.font.name = "Microsoft YaHei"; st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size); st.font.bold = True; st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True

# Running furniture
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(hp.add_run("REGENT · STRATEGY RESEARCH"), size=8.5, bold=True, color=MUTED)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(fp.add_run("Deep research · 2026-08-29"), size=8, color=MUTED)

# Cover/masthead
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(4)
set_font(p.add_run("深度调研"), size=10, bold=True, color=BLUE)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("Regent 突破方向判断"), size=25, bold=True, color=DARK)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18)
set_font(p.add_run("现有实现 × 开源生态 × AI 可靠性边界"), size=13, color=MUTED)
for label, value in [("决策问题", "哪个方向最适合形成首个可运行、可销售、可扩张的突破口？"),("建议方向", "Web/API 服务：异常到安全发布的闭环"),("范围", "截至 2026-08-29；技术、产品、AI 科学三方视角")]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(label + "："), size=10.5, bold=True, color=DARK); set_font(p.add_run(value), size=10.5)

call = doc.add_table(rows=1, cols=1); set_table_geometry(call, [9360]); shade(call.cell(0,0), LIGHT)
cp = call.cell(0,0).paragraphs[0]; cp.paragraph_format.space_before = Pt(7); cp.paragraph_format.space_after = Pt(7)
set_font(cp.add_run("核心结论  "), size=11, bold=True, color=BLUE)
set_font(cp.add_run("不做通用遗留接管或多端 App 托管；先把高信号线上问题稳定闭环到修复、验证、灰度和回滚。"), size=11, bold=True, color=DARK)

lines = SRC.read_text(encoding="utf-8").splitlines()
start = next(i for i,l in enumerate(lines) if l.startswith("## 执行结论"))
in_code = False; code_lines=[]; in_table=False; table_rows=[]

def flush_table():
    global table_rows
    if not table_rows: return
    rows = [r for r in table_rows if not all(re.fullmatch(r"[-: ]+", c or "") for c in r)]
    if len(rows) < 2: table_rows=[]; return
    cols = len(rows[0]); widths = [int(9360/cols)]*cols; widths[-1] += 9360-sum(widths)
    if cols == 6: widths=[1500,1100,1200,1200,1100,3260]
    t=doc.add_table(rows=len(rows), cols=cols); t.style="Table Grid"; set_table_geometry(t,widths)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            cell=t.cell(ri,ci); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri==0: shade(cell,LIGHT)
            p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.space_before=Pt(2)
            set_font(p.add_run(val.strip()), size=8.5 if cols>=5 else 9, bold=(ri==0), color=DARK)
    table_rows=[]

for raw in lines[start:]:
    line=raw.rstrip()
    if line.startswith("```"):
        if in_code:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
            r=p.add_run("\n".join(code_lines)); set_font(r,name="Consolas",size=8.5,color=DARK)
            in_code=False; code_lines=[]
        else:
            flush_table(); in_code=True
        continue
    if in_code: code_lines.append(line); continue
    if line.startswith("|") and line.endswith("|"):
        in_table=True; table_rows.append([c.strip() for c in line.strip("|").split("|")]); continue
    if in_table: flush_table(); in_table=False
    if not line: continue
    if line.startswith("## "): doc.add_paragraph(line[3:], style="Heading 1"); continue
    if line.startswith("### "): doc.add_paragraph(line[4:], style="Heading 2"); continue
    if line.startswith("> "):
        t=doc.add_table(rows=1,cols=1); set_table_geometry(t,[9360]); shade(t.cell(0,0),LIGHT)
        p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
        set_font(p.add_run(line[2:].replace("**","")),size=11,bold=True,color=DARK); continue
    m=re.match(r"^(\d+)\.\s+(.*)",line)
    if m:
        p=doc.add_paragraph(style="List Number"); p.paragraph_format.space_after=Pt(4)
        set_font(p.add_run(m.group(2).replace("**","")),size=10.5); continue
    if line.startswith("- "):
        content=line[2:]
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after=Pt(3)
        url_match=re.search(r"(https?://\S+)",content)
        if url_match:
            before=content[:url_match.start()].replace("**",""); url=url_match.group(1).rstrip(".,")
            set_font(p.add_run(before),size=9.5); add_hyperlink(p,url,url)
        else: set_font(p.add_run(content.replace("**","")),size=10.5)
        continue
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
    text=line.replace("**","").replace("`","")
    set_font(p.add_run(text),size=10.5)

flush_table()
doc.core_properties.title="Regent 突破方向深度调研"
doc.core_properties.subject="现有实现、开源生态与 AI 可靠性边界"
doc.core_properties.author="Regent Research"
doc.save(OUT)
print(OUT)
